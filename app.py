import streamlit as st
import fitz  # PyMuPDF pour traiter les PDF
import pytesseract  # Pour OCR (reconnaissance optique de caractères)
from PIL import Image  # Pour manipuler les images
from docx import Document  # Pour créer des documents Word
import io
import tempfile
import os

# Spécifiez le chemin d'accès à l'exécutable Tesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Modifiez ce chemin selon votre installation
# note: tesseract un est moteur de reconnaissance optique de caractere (OCR)


def pdf_to_images(pdf_path):
    """Convertit un PDF en une liste d'images PIL."""
    doc = fitz.open(pdf_path)
    images = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap()
        img = Image.open(io.BytesIO(pix.tobytes()))
        images.append(img)

    return images


def extract_text_from_images(images, progress_bar):
    """Extrait le texte de chaque image en utilisant Tesseract OCR."""
    text = ""
    num_images = len(images)

    for i, img in enumerate(images):
        # Convertir l'image en niveaux de gris
        gray_image = img.convert('L')

        # Utiliser pytesseract pour extraire le texte
        custom_config = r'--oem 3 --psm 6'
        text += pytesseract.image_to_string(gray_image, config=custom_config)

        # Mettre à jour la barre de progression
        progress_bar.progress((i + 1) / num_images)

    return text


def save_text_to_word(text, output_path):
    """Sauvegarde le texte extrait dans un document Word."""
    # Créer un document Word
    doc = Document()

    # Ajouter le texte extrait au document
    doc.add_paragraph(text)

    # Sauvegarder le document
    doc.save(output_path)


# Interface utilisateur avec Streamlit
st.title("PDF to Word Converter")

# Sélection du fichier PDF à convertir
uploaded_pdf = st.file_uploader("Choisissez un fichier PDF", type="pdf")

# Entrée pour spécifier l'emplacement de sauvegarde du fichier Word
output_path = st.text_input("Entrez le chemin d'enregistrement pour le fichier Word (par exemple, 'output.docx')")

# Bouton pour démarrer la conversion
if st.button("Convertir"):
    if uploaded_pdf is not None and output_path:
        # Écrit le fichier PDF téléchargé dans un fichier temporaire
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(uploaded_pdf.read())
            temp_pdf_path = temp_pdf.name

        # Convertir le PDF en images
        images = pdf_to_images(temp_pdf_path)

        # Barre de progression pour suivre l'avancement de l'extraction du texte
        progress_bar = st.progress(0)

        # Extraire le texte des images et mettre à jour la barre de progression
        text = extract_text_from_images(images, progress_bar)

        # Sauvegarder le texte extrait dans un fichier Word à l'emplacement spécifié
        save_text_to_word(text, output_path)

        # Indiquer que la conversion est terminée avec succès
        st.success(f"Conversion terminée! Le fichier a été sauvegardé à {output_path}")
    else:
        # Afficher une erreur si les entrées ne sont pas valides
        st.error("Veuillez télécharger un fichier PDF et entrer un chemin d'enregistrement valide.")



## methode 1
# import fitz  # PyMuPDF
# import pytesseract
# from PIL import Image
# from docx import Document
# import io
# import streamlit as st
#
# # Spécifiez le chemin d'accès à l'exécutable Tesseract
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Modifiez ce chemin selon votre installation
#
#
# def pdf_to_images(pdf_path):
#     doc = fitz.open(pdf_path)
#     images = []
#
#     for page_num in range(len(doc)):
#         page = doc.load_page(page_num)
#         pix = page.get_pixmap()
#         img = Image.open(io.BytesIO(pix.tobytes()))
#         images.append(img)
#
#     return images
#
#
# def extract_text_from_images(images):
#     text = ""
#
#     for img in images:
#         # Convertir l'image en niveaux de gris
#         gray_image = img.convert('L')
#
#         # Utiliser pytesseract pour extraire le texte
#         custom_config = r'--oem 3 --psm 6'
#         text += pytesseract.image_to_string(gray_image, config=custom_config)
#
#     return text
#
#
# def save_text_to_word(text, output_path):
#     # Créer un document Word
#     doc = Document()
#
#     # Ajouter le texte extrait au document
#     doc.add_paragraph(text)
#
#     # Sauvegarder le document
#     doc.save(output_path)
#
#
# def pdf_to_word(pdf_path, output_path):
#     images = pdf_to_images(pdf_path)
#     text = extract_text_from_images(images)
#     save_text_to_word(text, output_path)
#
#
# # Interface utilisateur Streamlit
# st.title("Convertisseur PDF vers Word")
# st.write(
#     "Téléchargez un fichier PDF et choisissez l'emplacement de sauvegarde pour convertir le PDF en document Word avec le texte extrait.")
#
# # Téléchargement du fichier PDF
# pdf_file = st.file_uploader("Choisissez un fichier PDF", type=["pdf"])
#
# # Sélection de l'emplacement de sauvegarde du fichier Word
# output_path = st.text_input(
#     "Entrez l'emplacement de sauvegarde du fichier Word (avec le nom du fichier, par exemple : output.docx)")
#
# # Bouton pour lancer le traitement
# if st.button("Convertir le PDF en Word") and pdf_file and output_path:
#     # Afficher une barre de progression
#     with st.spinner('Traitement en cours...'):
#         # Sauvegarder le fichier PDF téléchargé temporairement
#         with open("temp.pdf", "wb") as f:
#             f.write(pdf_file.read())
#
#         try:
#             # Conversion PDF en Word
#             pdf_to_word("temp.pdf", output_path)
#             st.success(f"Conversion terminée ! Le fichier a été sauvegardé à l'emplacement : {output_path}")
#         except Exception as e:
#             st.error(f"Erreur lors de la conversion : {e}")
#         finally:
#             # Supprimer le fichier PDF temporaire
#             import os
#
#             os.remove("temp.pdf")
